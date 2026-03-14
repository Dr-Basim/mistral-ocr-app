import streamlit as st
import base64
import re
import io
from mistralai import Mistral
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# يجب أن يكون set_page_config أول استدعاء لـ Streamlit دائماً
st.set_page_config(page_title="معالج الكتب الذكي", page_icon="📚", layout="wide")

# --- جلب مفتاح API ---
try:
    api_key = st.secrets["MISTRAL_API_KEY"]
except Exception:
    api_key = None

if api_key:
    client = Mistral(api_key=api_key)


# --- 1. تنظيف أولي بـ Regex ---
def basic_clean(text):
    # إزالة الرموز غير العربية غير المرغوبة مع الإبقاء على علامات الترقيم العربية والأرقام
    text = re.sub(r'[^\u0600-\u06FF\u0750-\u077F\s\.\،\؟\!\:\"\'\-\(\)\n0-9٠-٩]', ' ', text)
    text = re.sub(r' {2,}', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


# --- 2. تنظيف ذكي بـ LLM مع دعم النصوص الطويلة ---
def llm_clean_text(client, raw_text, progress_callback=None):
    CHUNK_WORDS = 2500
    words = raw_text.split()
    chunks = [' '.join(words[i:i+CHUNK_WORDS]) for i in range(0, len(words), CHUNK_WORDS)]
    total = len(chunks)
    cleaned_chunks = []

    system_prompt = (
        "أنت محرر نصوص عربية متخصص. مهمتك:\n"
        "1. إزالة الرموز والأحرف الغريبة الناتجة عن OCR\n"
        "2. تصحيح الكلمات العربية المكسورة أو المقطوعة\n"
        "3. إعادة تنظيم الفقرات بشكل طبيعي ومتسق\n"
        "4. الحفاظ على المحتوى الكامل دون حذف أو تلخيص\n"
        "5. لا تترجم ولا تعلّق، أعد النص النظيف فقط."
    )

    for idx, chunk in enumerate(chunks):
        if progress_callback:
            pct = 55 + int((idx / total) * 30)
            progress_callback(pct, f"🤖 التنظيف الذكي: جزء {idx+1} من {total}...")

        response = client.chat.complete(
            model="mistral-large-latest",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"نظّف هذا النص:\n\n{chunk}"}
            ]
        )
        cleaned_chunks.append(response.choices[0].message.content)

    return '\n\n'.join(cleaned_chunks)


# --- 3. إنشاء ملف Word مع دعم RTL العربي الكامل ---
def set_paragraph_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.insert(0, bidi)


def create_word_file(text):
    doc = Document()

    # ضبط اتجاه المستند RTL من إعدادات القسم
    section = doc.sections[0]
    sectPr = section._sectPr
    bidi_sect = OxmlElement('w:bidi')
    sectPr.insert(0, bidi_sect)

    # ضبط النمط الافتراضي
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = None

    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue

        # العناوين (أسطر تبدأ بـ #)
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            heading_text = line.lstrip('# ').strip()
            heading_style = f'Heading {min(level, 3)}'
            try:
                p = doc.add_heading(heading_text, level=min(level, 3))
            except Exception:
                p = doc.add_paragraph(heading_text)
        else:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Arial'
            # ضبط خط عربي للـ run
            rPr = run._r.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Arial')
            rFonts.set(qn('w:hAnsi'), 'Arial')
            rFonts.set(qn('w:cs'), 'Arial')
            rPr.insert(0, rFonts)

        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_rtl(p)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# --- 4. واجهة المستخدم ---
st.title("معالج الكتب العربي الذكي 🤖📖")
st.markdown("استخراج وتنظيف النصوص من كتب PDF المصوّرة باللغة العربية")
st.divider()

if not api_key:
    st.error("⚠️ لم يتم العثور على مفتاح MISTRAL_API_KEY في إعدادات Streamlit Secrets.")
    st.stop()

# خيارات المعالجة
col_opt1, col_opt2 = st.columns(2)
with col_opt1:
    use_llm_clean = st.checkbox(
        "✨ تفعيل التنظيف الذكي بالذكاء الاصطناعي",
        value=True,
        help="موصى به للكتب العربية المصوّرة — يُصحح الكلمات ويُنظم الفقرات"
    )
with col_opt2:
    show_page_headers = st.checkbox(
        "📄 إظهار أرقام الصفحات في الملف",
        value=False,
        help="إضافة فاصل --- صفحة X --- بين الصفحات"
    )

st.divider()
uploaded_file = st.file_uploader("⬆️ اختر ملف PDF", type=["pdf"])

if uploaded_file:
    file_size_mb = len(uploaded_file.getvalue()) / (1024 * 1024)
    st.info(f"📄 **{uploaded_file.name}** | الحجم: {file_size_mb:.1f} MB")

    if st.button("🚀 ابدأ عملية الاستخراج والتنسيق", type="primary", use_container_width=True):
        try:
            progress = st.progress(0, text="جاري التحضير...")

            # المرحلة 1: OCR
            progress.progress(10, text="🔄 يتم قراءة النص بـ Mistral OCR...")
            file_bytes = uploaded_file.read()
            encoded_pdf = base64.b64encode(file_bytes).decode("utf-8")

            ocr_response = client.ocr.process(
                model="mistral-ocr-latest",
                document={
                    "type": "document_url",
                    "document_url": f"data:application/pdf;base64,{encoded_pdf}"
                }
            )

            progress.progress(40, text="🧹 جاري التنظيف الأولي...")

            # تجميع النص من الصفحات
            raw_pages = []
            for i, page in enumerate(ocr_response.pages):
                if show_page_headers:
                    raw_pages.append(f"--- صفحة {i+1} ---\n{page.markdown}")
                else:
                    raw_pages.append(page.markdown)
            raw_text = '\n\n'.join(raw_pages)

            # تنظيف أولي
            cleaned_text = basic_clean(raw_text)

            # المرحلة 2: التنظيف الذكي
            if use_llm_clean:
                progress.progress(55, text="🤖 جاري التنظيف الذكي...")
                cleaned_text = llm_clean_text(client, cleaned_text, progress_callback=progress.progress)

            # المرحلة 3: إنشاء ملف Word
            progress.progress(90, text="📄 جاري إنشاء ملف Word...")
            word_data = create_word_file(cleaned_text)

            progress.progress(100, text="✅ اكتملت المعالجة!")
            st.success(f"🎉 تمت معالجة **{len(ocr_response.pages)} صفحة** بنجاح!")

            # معاينة النص
            with st.expander("👁️ معاينة النص المستخرج", expanded=False):
                st.text_area("", cleaned_text, height=400, label_visibility="collapsed")

            # أزرار التحميل
            st.subheader("⬇️ تحميل الملف")
            base_name = uploaded_file.name.replace('.pdf', '')
            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    label="📄 تحميل كملف Word (.docx)",
                    data=word_data,
                    file_name=f"{base_name}_extracted.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            with col2:
                st.download_button(
                    label="📝 تحميل كملف نصي (.txt)",
                    data=cleaned_text.encode('utf-8'),
                    file_name=f"{base_name}_extracted.txt",
                    mime="text/plain; charset=utf-8",
                    use_container_width=True
                )

        except Exception as e:
            st.error(f"❌ حدث خطأ: {e}")
            st.info("تأكد من أن مفتاح API صحيح وأن الملف بصيغة PDF صالحة.")

else:
    st.info("⬆️ يرجى رفع ملف PDF للبدء في الاستخراج.")
